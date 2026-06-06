"""遺伝のモデル実験 生徒用 記録用紙 (.docx) 生成スクリプト"""

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

JP_FONT = "ＭＳ ゴシック"
COLOR_PRIMARY = RGBColor(0x24, 0x6B, 0x46)
COLOR_BORDER = "CFDCD2"
COLOR_HEADER_FILL = "246B46"
COLOR_LIGHT_FILL = "EAF4EE"


def set_cell_shading(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, color="CFDCD2", sz="6"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), sz)
        b.set(qn("w:color"), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def style_run(run, *, size=10.5, bold=False, color=None, font=JP_FONT):
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)


def add_para(doc, text="", *, size=10.5, bold=False, color=None,
             align=None, space_before=0, space_after=4, indent=0):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if indent:
        pf.left_indent = Cm(indent)
    if text:
        run = p.add_run(text)
        style_run(run, size=size, bold=bold, color=color)
    return p


def add_heading_band(doc, text, *, size=13):
    """緑色の帯付き見出し"""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.rows[0].cells[0]
    cell.width = Cm(17)
    set_cell_shading(cell, COLOR_HEADER_FILL)
    set_cell_borders(cell, color=COLOR_HEADER_FILL, sz="0")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    style_run(run, size=size, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    add_para(doc, "", space_after=4)


def add_subheading(doc, text, *, size=11):
    p = add_para(doc, text, size=size, bold=True,
                 color=COLOR_PRIMARY, space_before=8, space_after=4)
    return p


def add_record_table(doc, n_rows, *, with_class=False):
    """記録テーブルを作る。with_class=True で末尾に集計列を増やす"""
    headers = ["回", "親A 遺伝子型", "親A の出した\nカード", "親B 遺伝子型",
               "親B の出した\nカード", "子の\n遺伝子型", "子の形質\n（丸／しわ）"]
    table = doc.add_table(rows=n_rows + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [Cm(1.0), Cm(2.6), Cm(2.6), Cm(2.6), Cm(2.6), Cm(2.4), Cm(2.6)]

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = widths[i]
        set_cell_shading(cell, COLOR_HEADER_FILL)
        set_cell_borders(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(h)
        style_run(run, size=9.5, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))

    for r in range(1, n_rows + 1):
        for c_i in range(len(headers)):
            cell = table.rows[r].cells[c_i]
            cell.width = widths[c_i]
            set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            if c_i == 0:
                run = p.add_run(str(r))
                style_run(run, size=10, color=RGBColor(0x6B, 0x7A, 0x73))


def add_count_table(doc, *, label="子の遺伝子型ごとの人数（自分の記録から）"):
    add_para(doc, label, size=10.5, bold=True, space_before=6, space_after=4)
    table = doc.add_table(rows=2, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    widths = [Cm(3.0), Cm(2.5), Cm(2.5), Cm(2.5), Cm(2.5)]
    headers = ["遺伝子型", "AA", "Aa", "aa", "合計"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = widths[i]
        set_cell_shading(cell, COLOR_LIGHT_FILL)
        set_cell_borders(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(h)
        style_run(run, size=10, bold=True, color=COLOR_PRIMARY)

    labels = ["人数（人）", "", "", "", ""]
    for i in range(5):
        cell = table.rows[1].cells[i]
        cell.width = widths[i]
        set_cell_borders(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        if i == 0:
            run = p.add_run(labels[i])
            style_run(run, size=10, bold=True)


def add_blank_lines(doc, n=3, *, leader="／"):
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run("　" + "＿" * 50)
        style_run(run, size=10.5, color=RGBColor(0xCF, 0xDC, 0xD2))


def add_info_header(doc):
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [Cm(2.0), Cm(5.0), Cm(2.0), Cm(8.0)]
    labels = ["クラス", "", "氏名", ""]
    for i, label in enumerate(labels):
        cell = table.rows[0].cells[i]
        cell.width = widths[i]
        set_cell_borders(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        if i % 2 == 0:
            set_cell_shading(cell, COLOR_LIGHT_FILL)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(label)
            style_run(run, size=10.5, bold=True, color=COLOR_PRIMARY)

    table2 = doc.add_table(rows=1, cols=4)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths2 = [Cm(2.0), Cm(3.0), Cm(2.0), Cm(10.0)]
    labels2 = ["出席番号", "", "実施日", ""]
    for i, label in enumerate(labels2):
        cell = table2.rows[0].cells[i]
        cell.width = widths2[i]
        set_cell_borders(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        if i % 2 == 0:
            set_cell_shading(cell, COLOR_LIGHT_FILL)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(label)
            style_run(run, size=10.5, bold=True, color=COLOR_PRIMARY)


def add_class_summary_table(doc):
    add_para(doc, "クラス全体の集計", size=11, bold=True,
             color=COLOR_PRIMARY, space_before=10, space_after=4)
    add_para(doc, "アプリ画面右の「クラス集計」を見て、クラス全体の人数を書きうつしましょう。",
             size=9.5, color=RGBColor(0x4F, 0x60, 0x56), space_after=4)
    table = doc.add_table(rows=4, cols=4)
    widths = [Cm(4.0), Cm(3.0), Cm(3.0), Cm(3.0)]
    rows = [
        ["", "AA", "Aa", "aa"],
        ["クラス全体の人数", "", "", ""],
        ["遺伝子型の比", "", "", ""],
        ["", "丸", "しわ", ""],
    ]
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = table.rows[r_i].cells[c_i]
            cell.width = widths[c_i]
            set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            is_header = (r_i == 0) or (c_i == 0 and r_i in (1, 2))
            if is_header or r_i == 3:
                set_cell_shading(cell, COLOR_LIGHT_FILL)
            run = p.add_run(val)
            style_run(run, size=10, bold=is_header,
                      color=COLOR_PRIMARY if is_header else None)

    add_para(doc, "形質の比　丸：しわ ＝ 　　　　　　：　　　　　　",
             size=10.5, space_before=6, space_after=4)


def main():
    doc = Document()

    section = doc.sections[0]
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)

    style = doc.styles["Normal"]
    style.font.name = JP_FONT
    style.font.size = Pt(10.5)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), JP_FONT)

    # ----- タイトル -----
    title = add_para(doc, "探Q実習1　遺伝のモデル実験　記録用紙",
                     size=18, bold=True, color=COLOR_PRIMARY,
                     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_para(doc, "中学校3年 理科　「遺伝の規則性と遺伝子」",
             size=10, color=RGBColor(0x4F, 0x60, 0x56),
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

    add_info_header(doc)
    add_para(doc, "", space_after=2)

    # ----- 学習のめあて -----
    add_heading_band(doc, "■ 学習のめあて")
    add_para(doc,
             "純系どうしや雑種どうしの親から子・孫の遺伝子の組み合わせをモデルでつくり、"
             "その結果を集計することで、遺伝の規則性（メンデルの法則）を確かめる。",
             size=10.5, space_after=8)

    # ----- 課題 -----
    add_heading_band(doc, "■ 探究の課題（自分の予想）")
    add_para(doc, "課題：純系の親（AA × aa）からは、どんな子が生まれるだろうか？",
             size=10.5, bold=True, space_after=2)
    add_para(doc, "予想（自分の言葉で）：",
             size=10.5, space_after=2)
    add_blank_lines(doc, 2)
    add_para(doc, "課題：雑種の親（Aa × Aa）からは、どんな比で孫が生まれるだろうか？",
             size=10.5, bold=True, space_before=4, space_after=2)
    add_para(doc, "予想（自分の言葉で）：",
             size=10.5, space_after=2)
    add_blank_lines(doc, 2)

    # ----- ステップ1 -----
    doc.add_page_break()
    add_heading_band(doc, "■ ステップ1　親（AA × aa）→ 子をつくる")
    add_para(doc,
             "先生の合図で初期化されたあと、相手を変えて 5回 ペア交配しよう。"
             "親役は AA か aa のどちらかになっています。",
             size=10, space_after=6)
    add_record_table(doc, n_rows=5)
    add_para(doc, "", space_after=4)
    add_subheading(doc, "ステップ1のまとめ")
    add_para(doc, "・子の遺伝子型は、すべて 　　　　　　　　　 だった。",
             size=10.5, space_after=4)
    add_para(doc, "・子の形質（見た目）は、すべて 　　　　　　　　　 だった。",
             size=10.5, space_after=4)
    add_para(doc, "・このことから言えること：",
             size=10.5, space_after=2)
    add_blank_lines(doc, 3)

    # ----- ステップ2 -----
    doc.add_page_break()
    add_heading_band(doc, "■ ステップ2　子（Aa × Aa）→ 孫をつくる")
    add_para(doc,
             "先生の合図で全員が Aa（雑種）に初期化されたあと、相手を変えて 40回 ペア交配しよう。"
             "親役は両方とも Aa です。",
             size=10, space_after=6)

    add_subheading(doc, "前半（1〜20回）")
    add_record_table(doc, n_rows=20)
    doc.add_page_break()
    add_subheading(doc, "後半（21〜40回）")
    add_record_table(doc, n_rows=20)

    add_para(doc, "", space_after=6)
    add_count_table(doc, label="自分が記録した40回ぶんの人数")

    add_para(doc, "自分の記録での比　AA : Aa : aa ＝ 　　　：　　　：　　　",
             size=10.5, space_before=6, space_after=4)
    add_para(doc, "自分の記録での比　丸 : しわ ＝ 　　　：　　　",
             size=10.5, space_after=8)

    # ----- クラス全体集計 -----
    doc.add_page_break()
    add_heading_band(doc, "■ クラス全体の集計と考察")
    add_class_summary_table(doc)

    add_subheading(doc, "考察1　予想とのちがいはあったか？")
    add_blank_lines(doc, 4)

    add_subheading(doc, "考察2　なぜそのような比になったのか？")
    add_para(doc,
             "（ヒント：Aa × Aa の親が出すカードの組み合わせを表にして考えよう）",
             size=9.5, color=RGBColor(0x4F, 0x60, 0x56), space_after=4)
    # ヒント表（パネット表）
    table = doc.add_table(rows=3, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    panels = [
        ["", "親B：A", "親B：a"],
        ["親A：A", "", ""],
        ["親A：a", "", ""],
    ]
    for r_i, row in enumerate(panels):
        for c_i, val in enumerate(row):
            cell = table.rows[r_i].cells[c_i]
            cell.width = Cm(3.5)
            set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            if r_i == 0 or c_i == 0:
                set_cell_shading(cell, COLOR_LIGHT_FILL)
            run = p.add_run(val)
            style_run(run, size=10.5, bold=(r_i == 0 or c_i == 0),
                      color=COLOR_PRIMARY if (r_i == 0 or c_i == 0) else None)

    add_para(doc, "", space_after=4)
    add_blank_lines(doc, 4)

    add_subheading(doc, "考察3　もっと回数を増やしたら、比はどうなると思うか？")
    add_blank_lines(doc, 3)

    # ----- まとめ -----
    add_heading_band(doc, "■ 学んだこと・気づいたこと")
    add_blank_lines(doc, 5)

    out = "遺伝のモデル実験_記録用紙.docx"
    doc.save(out)
    print(f"Generated: {out}")


if __name__ == "__main__":
    main()
